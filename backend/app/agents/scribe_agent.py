import logging
from typing import Dict, Any, List
from app.agents.base_agent import BaseAgent
from app.services.artifact_service import artifact_service
import docx
import io

logger = logging.getLogger(__name__)

class ScribeAgent(BaseAgent):
    """Agent responsible for requirements analysis and document generation."""
    
    async def run(self, context: Dict[str, Any]) -> Dict[str, Any]:
        self.logger.info("SCRIBE: Starting document generation...")
        
        scribe_config = context.get("scribe", {})
        selected_docs = scribe_config.get("selected_documents", ["feature_doc"])
        user_prompt = context.get("user_prompt", "")
        requirement_text = scribe_config.get("requirement_text", "")
        project_context = scribe_config.get("project_context", "")
        output_format = scribe_config.get("output_format", "markdown")

        results = {}
        generated_contents = {}
        
        for doc_type in selected_docs:
            prompt = f"""
Generate a {doc_type.replace('_', ' ').upper()} based on the following:
REQUIREMENT: {requirement_text}
PROJECT CONTEXT: {project_context}
USER NOTES: {user_prompt}

The output should be high-quality markdown.
"""
            # Call LLM
            content = await self.call_llm(prompt)
            generated_contents[doc_type] = content
            
            # Save artifact
            if output_format == "docx":
                # Convert Markdown to DOCX
                doc = docx.Document()
                doc.add_heading(doc_type.replace('_', ' ').upper(), 0)

                # Simple markdown parsing (headings and paragraphs)
                for line in content.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('# '):
                        doc.add_heading(line[2:], 1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], 2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], 3)
                    elif line.startswith('- ') or line.startswith('* '):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line)

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                docx_bytes = bio.read()

                filename = f"{doc_type}.docx"
                path = artifact_service.save_artifact(self.task_id, doc_type, docx_bytes, filename=filename)
                results[doc_type] = path
                self.logger.info(f"SCRIBE: Generated {doc_type} (docx) and saved to {path}")
            else:
                path = artifact_service.save_artifact(self.task_id, doc_type, content)
                results[doc_type] = path
                self.logger.info(f"SCRIBE: Generated {doc_type} and saved to {path}")

        # Validate outputs against guardrails
        for doc_type, content_text in generated_contents.items():
            if not self.validate_output(content_text):
                self.logger.warning(f"SCRIBE: Output validation failed for {doc_type}")

        artifact_paths = list(results.values())

        return {
            "status": "success",
            "message": f"Generated {len(results)} documents",
            "artifacts": results,
            "artifact_paths": artifact_paths,
            "summary": f"Generated {', '.join(selected_docs)} documents"
        }
